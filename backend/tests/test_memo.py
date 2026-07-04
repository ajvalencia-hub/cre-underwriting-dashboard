"""F7: IC memo generation — docx round-trip on section presence, rendered
numbers matching the scenario outputs exactly, graceful degradation of
optional sections, and the route wire-up (fresh compute path + quickscreen
rejection)."""

import io
import json
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.database import Base, get_db
from app.main import app
from app.services import memo_service
from app.services.proforma import engine

FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture
def analytic() -> dict:
    return json.loads((FIXTURES / "analytic_acquisition.json").read_text())


def _doc_text(memo_bytes: bytes) -> str:
    doc = Document(io.BytesIO(memo_bytes))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def _full_memo_bytes(analytic) -> bytes:
    computed = engine.compute(analytic)
    return memo_service.build_memo(
        deal_name="Maple Court",
        scenario_name="Base Case",
        inputs=analytic,
        outputs=computed["outputs"],
        debt=computed["debt"],
        sources_and_uses=computed["sourcesAndUses"],
        sensitivity={
            "description": "Rent x exit cap",
            "header": ["", "-25bps", "base", "+25bps"],
            "rows": [["-5% rent", "10.1%", "10.9%", "11.5%"]],
        },
        benchmark_flags=[
            {"verdict": "warning", "explanation": "Rent above the 85th percentile.",
             "source": "census_acs + hud", "asOf": "2022"},
        ],
    )


def test_memo_contains_every_section(analytic):
    text = _doc_text(_full_memo_bytes(analytic))
    for section in (
        "EXECUTIVE SUMMARY", "SOURCES & USES", "KEY ASSUMPTIONS", "RETURNS SUMMARY",
        "DEBT SUMMARY", "SENSITIVITY", "MARKET CONTEXT", "ASSUMPTIONS & LIMITATIONS",
    ):
        assert section in text, f"missing memo section: {section}"
    assert "Maple Court" in text
    assert "Rent above the 85th percentile." in text


def test_rendered_numbers_match_the_computed_outputs(analytic):
    computed = engine.compute(analytic)["outputs"]
    text = _doc_text(_full_memo_bytes(analytic))
    # Three metrics parsed back out of the document and checked against the
    # engine's numbers via the exact same formatting rules.
    assert memo_service.format_value(computed["equityMultiple"], "multiple") == "1.55x"
    assert "1.55x" in text
    assert memo_service.format_value(computed["terminalValue"], "currency") == "$1,000,000"
    assert "$1,000,000" in text
    assert memo_service.format_value(computed["leveredIrr"], "percent") == "11.57%"
    assert "11.57%" in text


def test_optional_sections_degrade_cleanly(analytic):
    memo_bytes = memo_service.build_memo(
        deal_name="Bare Deal",
        scenario_name="Stored Only",
        inputs={"dealType": "acquisition", "holdPeriodYears": 5},
        outputs={"equityMultiple": 1.42},
        debt=None,
        sources_and_uses=None,
        sensitivity=None,
        benchmark_flags=None,
    )
    text = _doc_text(memo_bytes)
    assert "RETURNS SUMMARY" in text
    assert "1.42x" in text
    for absent in ("DEBT SUMMARY", "SENSITIVITY", "MARKET CONTEXT", "SOURCES & USES"):
        assert absent not in text
    assert "ASSUMPTIONS & LIMITATIONS" in text  # boilerplate always present


def test_custom_limitations_text_replaces_boilerplate(analytic):
    memo_bytes = memo_service.build_memo(
        deal_name="X", scenario_name="Y", inputs={}, outputs={},
        limitations_text="Custom internal-use disclaimer.",
    )
    text = _doc_text(memo_bytes)
    assert "Custom internal-use disclaimer." in text
    assert memo_service.DEFAULT_LIMITATIONS[:40] not in text


# ------------------------------------------------------------------ route

@pytest.fixture
def client(monkeypatch):
    sql_engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    Base.metadata.create_all(sql_engine)
    TestSession = sessionmaker(bind=sql_engine)

    def _override():
        db = TestSession()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = _override
    # The memo route consults benchmarks when an address/market exists — keep
    # tests offline.
    from app.routers import scenarios as scenarios_router

    monkeypatch.setattr(
        scenarios_router.benchmarks, "build_benchmarks",
        lambda *a, **k: {"location": {}, "flags": [], "unavailable": []},
    )
    yield TestClient(app)
    app.dependency_overrides.pop(get_db)
    sql_engine.dispose()


def test_memo_route_returns_docx_via_fresh_compute(client, analytic):
    # Since the native engine landed, full scenarios don't need a template —
    # they create directly and the memo computes fresh from their inputs.
    created = client.post(
        "/api/scenarios",
        json={"scenarioName": "Base", "kind": "full", "templateId": None,
              "mappingProfileId": None, "inputs": analytic},
    )
    assert created.status_code == 200

    resp = client.post(f"/api/scenarios/{created.json()['id']}/memo", json={})
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml"
    )
    text = _doc_text(resp.content)
    assert "RETURNS SUMMARY" in text
    assert "1.55x" in text  # fresh engine compute of the analytic fixture


def test_memo_route_rejects_quickscreen(client):
    created = client.post(
        "/api/scenarios",
        json={"scenarioName": "QS", "kind": "quickscreen", "inputs": {"rent": 1800}},
    ).json()
    resp = client.post(f"/api/scenarios/{created['id']}/memo", json={})
    assert resp.status_code == 400


def test_memo_route_422_when_no_outputs_and_insufficient_inputs(client):
    from app.models import Scenario

    db = next(app.dependency_overrides[get_db]())
    row = Scenario(scenario_name="Empty", kind="full", inputs={"dealType": "acquisition"}, outputs={})
    db.add(row)
    db.commit()
    db.refresh(row)

    resp = client.post(f"/api/scenarios/{row.id}/memo", json={})
    assert resp.status_code == 422
    assert "purchasePrice" in resp.json()["detail"]
