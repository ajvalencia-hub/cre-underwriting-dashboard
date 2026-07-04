"""IC memo generation via python-docx.

STRICT RULE: this module contains ZERO financial math. Every number is a
formatted pass-through from stored scenario outputs or a fresh engine
compute handed in by the router — nothing is derived, summed, or
recalculated here. Formatting only: $#,##0 / 0.00% / 0.00x per the schema
output type.
"""

import json
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.config import FIRM_NAME, MEMO_BRAND_COLOR

_SCHEMA_PATH = Path(__file__).resolve().parent.parent / "data" / "input_schema.json"
_OUTPUT_META: dict[str, dict] = {
    o["id"]: o for o in json.loads(_SCHEMA_PATH.read_text())["outputs"]
}

DEFAULT_LIMITATIONS = (
    "This memorandum was prepared for internal investment-committee discussion "
    "only. Projections are estimates based on the stated assumptions; actual "
    "results will differ. Market benchmark data comes from public sources "
    "(Census ACS, HUD, FHFA, BLS, FEMA) as of the dates noted and has not been "
    "independently verified. This document is not an offer to sell or a "
    "solicitation of an offer to buy any security."
)

_RETURNS_MEMO_IDS = [
    "unleveredIrr", "leveredIrr", "lpIrr", "gpIrr", "equityMultiple",
    "lpEquityMultiple", "cashOnCashYear1", "stabilizedCashOnCash",
    "paybackPeriodYears", "npv", "yieldOnCost", "goingInCapRate",
    "developmentSpreadBps", "terminalValue", "netSaleProceeds", "totalProfit",
]
_KEY_ASSUMPTION_FIELDS = [
    ("dealType", "Deal type", "text"),
    ("propertyType", "Property type", "text"),
    ("purchasePrice", "Purchase price", "currency"),
    ("landCost", "Land cost", "currency"),
    ("hardCosts", "Hard costs", "currency"),
    ("grossPotentialRent", "Gross potential rent", "currency"),
    ("vacancyPct", "Vacancy", "percent"),
    ("rentGrowthPct", "Rent growth", "percent"),
    ("expenseGrowthPct", "Expense growth", "percent"),
    ("holdPeriodYears", "Hold period (yrs)", "number"),
    ("exitCapRatePct", "Exit cap rate", "percent"),
    ("interestRate", "Interest rate", "percent"),
    ("ltvOrLtc", "LTV / LTC", "percent"),
    ("amortYears", "Amortization (yrs)", "number"),
    ("ioMonths", "IO period (mo)", "number"),
]


def format_value(value, value_type: str) -> str:
    if value is None:
        return "—"
    if value_type == "currency":
        return f"${value:,.0f}"
    if value_type == "percent":
        return f"{value * 100:.2f}%"
    if value_type == "multiple":
        return f"{value:.2f}x"
    if value_type == "years":
        return f"{value:.1f} yrs"
    if value_type == "number":
        return f"{value:,.2f}".rstrip("0").rstrip(".")
    return str(value)


def _brand_rgb() -> RGBColor:
    return RGBColor.from_string(MEMO_BRAND_COLOR)


def _heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _brand_rgb()
    paragraph.space_after = Pt(4)


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        cell = table.rows[i].cells[1]
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _grid_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for c, text in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            table.rows[r].cells[c].text = text


def _numbered_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_memo(
    deal_name: str,
    scenario_name: str,
    inputs: dict,
    outputs: dict,
    debt: dict | None = None,
    sources_and_uses: dict | None = None,
    sensitivity: dict | None = None,
    benchmark_flags: list[dict] | None = None,
    limitations_text: str | None = None,
    conventions: dict | None = None,
) -> bytes:
    """Assemble the .docx and return its bytes. Optional sections (debt,
    sources & uses, sensitivity, benchmarks) are silently omitted when their
    data isn't provided — never fabricated."""
    doc = Document()

    # Header: firm + deal + date.
    header_paragraph = doc.sections[0].header.paragraphs[0]
    header_paragraph.text = f"{FIRM_NAME} — {deal_name} — {date.today().isoformat()}"
    _numbered_footer(doc)

    title = doc.add_paragraph()
    run = title.add_run(f"Investment Committee Memorandum — {deal_name}")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = _brand_rgb()
    doc.add_paragraph(f"Scenario: {scenario_name} · Prepared {date.today():%B %d, %Y}")

    # ---- Executive summary ------------------------------------------------
    _heading(doc, "Executive Summary")
    summary_bits = []
    for output_id, label in (
        ("leveredIrr", "levered IRR"),
        ("equityMultiple", "equity multiple"),
        ("yieldOnCost", "yield on cost"),
        ("minDscr", "min DSCR"),
    ):
        value = outputs.get(output_id)
        if value is not None and not isinstance(value, str):
            summary_bits.append(
                f"{format_value(value, _OUTPUT_META.get(output_id, {}).get('type', 'number'))} {label}"
            )
    deal_type = inputs.get("dealType", "deal")
    hold = inputs.get("holdPeriodYears")
    doc.add_paragraph(
        f"{deal_name} is underwritten as a {deal_type} with a "
        f"{format_value(hold, 'number') if hold else '—'}-year hold"
        + (f", projecting {', '.join(summary_bits)}." if summary_bits else ".")
    )

    # ---- Sources & uses ----------------------------------------------------
    if sources_and_uses and (sources_and_uses.get("uses") or sources_and_uses.get("sources")):
        _heading(doc, "Sources & Uses")
        rows = [
            (label, format_value(amount, "currency"))
            for label, amount in sources_and_uses.get("uses", [])
            if amount
        ]
        if rows:
            doc.add_paragraph("Uses").runs[0].font.bold = True
            _kv_table(doc, rows)
        source_rows = [
            (label, format_value(amount, "currency"))
            for label, amount in sources_and_uses.get("sources", [])
            if amount
        ]
        if source_rows:
            doc.add_paragraph("Sources").runs[0].font.bold = True
            _kv_table(doc, source_rows)

    # ---- Key assumptions ---------------------------------------------------
    _heading(doc, "Key Assumptions")
    assumption_rows = [
        (label, format_value(inputs[field_id], value_type))
        for field_id, label, value_type in _KEY_ASSUMPTION_FIELDS
        if inputs.get(field_id) not in (None, "", 0)
    ]
    if conventions:
        if conventions.get("irrConvention"):
            assumption_rows.append(
                (
                    "IRR convention",
                    "Date-based XIRR (Actual/365)"
                    if conventions["irrConvention"] == "xirr"
                    else "Periodic monthly, annualized",
                )
            )
        if conventions.get("waterfallStyle"):
            assumption_rows.append(
                (
                    "Waterfall style",
                    "American (deal-by-deal ledger)"
                    if conventions["waterfallStyle"] == "american"
                    else "European (whole-fund, IRR hurdles)",
                )
            )
    if assumption_rows:
        _kv_table(doc, assumption_rows)

    # ---- Returns summary ----------------------------------------------------
    _heading(doc, "Returns Summary")
    return_rows = [
        (
            _OUTPUT_META[output_id]["label"],
            format_value(outputs[output_id], _OUTPUT_META[output_id]["type"]),
        )
        for output_id in _RETURNS_MEMO_IDS
        if output_id in outputs and not isinstance(outputs[output_id], str)
    ]
    if return_rows:
        _kv_table(doc, return_rows)
    else:
        doc.add_paragraph("No computed return metrics are stored for this scenario.")

    # ---- Debt summary --------------------------------------------------------
    if debt:
        _heading(doc, "Debt Summary")
        _kv_table(
            doc,
            [
                ("Loan amount", format_value(debt.get("loanAmount"), "currency")),
                ("Constraint-sized proceeds", format_value(debt.get("sizedLoanAmount"), "currency")),
                ("Governing constraint", str(debt.get("governingConstraint", "—"))),
                ("Sizing NOI", format_value(debt.get("sizingNoi"), "currency")),
            ],
        )
        stress = debt.get("stress") or []
        if stress:
            doc.add_paragraph("Rate / NOI stress").runs[0].font.bold = True
            _grid_table(
                doc,
                ["Scenario", "DSCR", "Refi proceeds", "Shortfall"],
                [
                    [
                        "Base" if c["rateBumpBps"] == 0 and c["noiHaircutPct"] == 0
                        else " · ".join(
                            part for part in (
                                f"+{c['rateBumpBps']}bps" if c["rateBumpBps"] else None,
                                f"NOI −{round(c['noiHaircutPct'] * 100)}%" if c["noiHaircutPct"] else None,
                            ) if part
                        ),
                        format_value(c.get("dscr"), "multiple"),
                        format_value(c.get("refiProceeds"), "currency"),
                        format_value(c.get("refiShortfall"), "currency") if c.get("refiShortfall") else "—",
                    ]
                    for c in stress
                ],
            )

    # ---- Sensitivity matrix ---------------------------------------------------
    if sensitivity and sensitivity.get("rows"):
        _heading(doc, "Sensitivity")
        if sensitivity.get("description"):
            doc.add_paragraph(str(sensitivity["description"]))
        _grid_table(
            doc,
            [str(h) for h in sensitivity.get("header", [])],
            [[str(cell) for cell in row] for row in sensitivity["rows"]],
        )

    # ---- Market context flags ---------------------------------------------------
    if benchmark_flags:
        _heading(doc, "Market Context")
        for flag in benchmark_flags:
            marker = {"ok": "✓", "caution": "△", "warning": "✕"}.get(flag.get("verdict", ""), "•")
            source = flag.get("source", "")
            as_of = flag.get("asOf", "")
            provenance = f" [{source}{', ' + as_of if as_of else ''}]" if source else ""
            doc.add_paragraph(f"{marker} {flag.get('explanation', '')}{provenance}")

    # ---- Assumptions & limitations -------------------------------------------
    _heading(doc, "Assumptions & Limitations")
    doc.add_paragraph(limitations_text or DEFAULT_LIMITATIONS)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
